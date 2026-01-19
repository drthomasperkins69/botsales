"""
Markdown to Word document converter.
"""

import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


class MarkdownToWordConverter:
    """Converts markdown content to Word document format."""

    def __init__(self):
        self.document = None

    def convert(self, markdown_files, output_path):
        """
        Convert multiple markdown files to a single Word document.

        Args:
            markdown_files: List of dicts with 'filename' and 'content' keys
            output_path: Path to save the output .docx file
        """
        self.document = Document()
        self._setup_styles()

        for i, md_file in enumerate(markdown_files):
            # Add page break between files (except for the first one)
            if i > 0:
                self.document.add_page_break()

            # Add filename as a section header
            self._add_section_header(md_file['filename'])

            # Convert and add the markdown content
            self._convert_markdown(md_file['content'])

        self.document.save(output_path)

    def _setup_styles(self):
        """Set up document styles."""
        styles = self.document.styles

        # Modify the Normal style
        normal_style = styles['Normal']
        normal_style.font.name = 'Calibri'
        normal_style.font.size = Pt(11)

    def _add_section_header(self, filename):
        """Add a section header showing the source filename."""
        # Remove .md extension for cleaner display
        display_name = filename.rsplit('.', 1)[0] if '.' in filename else filename

        para = self.document.add_paragraph()
        run = para.add_run(f"Source: {display_name}")
        run.italic = True
        run.font.size = Pt(10)
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Add a horizontal line
        self.document.add_paragraph('─' * 50)

    def _convert_markdown(self, content):
        """Convert markdown content to Word document elements."""
        lines = content.split('\n')
        i = 0
        in_code_block = False
        code_block_content = []

        while i < len(lines):
            line = lines[i]

            # Handle code blocks
            if line.strip().startswith('```'):
                if in_code_block:
                    # End of code block
                    self._add_code_block('\n'.join(code_block_content))
                    code_block_content = []
                    in_code_block = False
                else:
                    # Start of code block
                    in_code_block = True
                i += 1
                continue

            if in_code_block:
                code_block_content.append(line)
                i += 1
                continue

            # Handle headers
            if line.startswith('#'):
                self._add_header(line)
                i += 1
                continue

            # Handle horizontal rules
            if re.match(r'^(\*{3,}|-{3,}|_{3,})\s*$', line.strip()):
                self.document.add_paragraph('─' * 50)
                i += 1
                continue

            # Handle unordered lists
            if re.match(r'^\s*[-*+]\s+', line):
                list_items = []
                while i < len(lines) and re.match(r'^\s*[-*+]\s+', lines[i]):
                    list_items.append(lines[i])
                    i += 1
                self._add_unordered_list(list_items)
                continue

            # Handle ordered lists
            if re.match(r'^\s*\d+\.\s+', line):
                list_items = []
                while i < len(lines) and re.match(r'^\s*\d+\.\s+', lines[i]):
                    list_items.append(lines[i])
                    i += 1
                self._add_ordered_list(list_items)
                continue

            # Handle blockquotes
            if line.startswith('>'):
                quote_lines = []
                while i < len(lines) and lines[i].startswith('>'):
                    quote_lines.append(lines[i])
                    i += 1
                self._add_blockquote(quote_lines)
                continue

            # Handle regular paragraphs
            if line.strip():
                # Collect paragraph lines until empty line
                para_lines = []
                while i < len(lines) and lines[i].strip() and not lines[i].startswith('#'):
                    if re.match(r'^\s*[-*+]\s+', lines[i]) or re.match(r'^\s*\d+\.\s+', lines[i]):
                        break
                    if lines[i].startswith('>'):
                        break
                    if lines[i].strip().startswith('```'):
                        break
                    para_lines.append(lines[i])
                    i += 1
                self._add_paragraph(' '.join(para_lines))
                continue

            i += 1

    def _add_header(self, line):
        """Add a header to the document."""
        # Count the number of # symbols
        match = re.match(r'^(#{1,6})\s+(.+)$', line)
        if match:
            level = len(match.group(1))
            text = match.group(2).strip()

            # Map markdown levels to Word heading styles
            heading_style = f'Heading {min(level, 9)}'

            try:
                para = self.document.add_paragraph(text, style=heading_style)
            except KeyError:
                # Fallback if heading style doesn't exist
                para = self.document.add_paragraph()
                run = para.add_run(text)
                run.bold = True
                run.font.size = Pt(16 - level)

    def _add_paragraph(self, text):
        """Add a paragraph with inline formatting."""
        para = self.document.add_paragraph()
        self._add_formatted_text(para, text)

    def _add_formatted_text(self, para, text):
        """Add text to a paragraph with inline formatting (bold, italic, code)."""
        # Pattern to match inline formatting
        pattern = r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|__(.+?)__|_(.+?)_|`(.+?)`|\[(.+?)\]\((.+?)\))'

        last_end = 0
        for match in re.finditer(pattern, text):
            # Add text before the match
            if match.start() > last_end:
                para.add_run(text[last_end:match.start()])

            full_match = match.group(0)

            if match.group(2):  # Bold and italic ***text***
                run = para.add_run(match.group(2))
                run.bold = True
                run.italic = True
            elif match.group(3):  # Bold **text**
                run = para.add_run(match.group(3))
                run.bold = True
            elif match.group(4):  # Italic *text*
                run = para.add_run(match.group(4))
                run.italic = True
            elif match.group(5):  # Bold __text__
                run = para.add_run(match.group(5))
                run.bold = True
            elif match.group(6):  # Italic _text_
                run = para.add_run(match.group(6))
                run.italic = True
            elif match.group(7):  # Inline code `text`
                run = para.add_run(match.group(7))
                run.font.name = 'Courier New'
            elif match.group(8) and match.group(9):  # Link [text](url)
                run = para.add_run(f"{match.group(8)} ({match.group(9)})")
                run.underline = True

            last_end = match.end()

        # Add any remaining text
        if last_end < len(text):
            para.add_run(text[last_end:])

    def _add_code_block(self, code):
        """Add a code block to the document."""
        para = self.document.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.5)
        run = para.add_run(code)
        run.font.name = 'Courier New'
        run.font.size = Pt(10)

    def _add_unordered_list(self, items):
        """Add an unordered list to the document."""
        for item in items:
            # Extract the text after the bullet
            text = re.sub(r'^\s*[-*+]\s+', '', item)
            para = self.document.add_paragraph(style='List Bullet')
            self._add_formatted_text(para, text)

    def _add_ordered_list(self, items):
        """Add an ordered list to the document."""
        for item in items:
            # Extract the text after the number
            text = re.sub(r'^\s*\d+\.\s+', '', item)
            para = self.document.add_paragraph(style='List Number')
            self._add_formatted_text(para, text)

    def _add_blockquote(self, lines):
        """Add a blockquote to the document."""
        # Combine lines and remove > prefix
        text_lines = []
        for line in lines:
            text_lines.append(re.sub(r'^>\s*', '', line))
        text = ' '.join(text_lines)

        para = self.document.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.5)
        para.paragraph_format.right_indent = Inches(0.5)
        run = para.add_run(text)
        run.italic = True
